import streamlit as st
from modules.logger import get_logger
# Configuração inicial do Streamlit
st.set_page_config(page_title="Agente de IA para Professores", layout="wide")

import os
from langchain.schema import SystemMessage, HumanMessage, AIMessage
from langchain.chat_models import ChatOpenAI
from modules.chat import Chatbot
from modules.file_loader import process_uploaded_files
from modules.vector_store import create_vector_store
from docx import Document
from io import BytesIO

# Inicializar logger
logger = get_logger(__name__)

# Log de sucesso ou erro
try:
    vector_store = create_vector_store(documents, api_key)
    logger.info("Armazenamento vetorial criado com sucesso.")
except Exception as e:
    logger.error("Erro ao criar o armazenamento vetorial", exc_info=True)
    st.error("Erro ao criar armazenamento vetorial. Consulte o log para mais detalhes.")
    
def export_to_docx(lesson_plan, file_name="Plano_de_Aula.docx"):
    """Exporta o plano de aula para um arquivo DOCX."""
    document = Document()
    document.add_heading("Plano de Aula", level=1)
    
    # Adiciona o conteúdo do plano de aula
    for line in lesson_plan.split("\n"):
        document.add_paragraph(line)
    
    # Salva o arquivo em memória
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

# Configuração inicial
st.set_page_config(page_title="Agente de IA para Professores", layout="wide")
st.title("Agente de IA para Professores")

# Configurações de API
try:
    api_key = st.secrets["OPENAI_API_KEY"]
except KeyError:
    st.error("Por favor, configure a variável OPENAI_API_KEY nos segredos do Streamlit.")
    st.stop()

# Upload de arquivos
st.sidebar.header("Upload de Arquivos")
uploaded_files = st.sidebar.file_uploader(
    "Carregue arquivos (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True
)

documents = []
if uploaded_files:
    try:
        documents = process_uploaded_files(uploaded_files)
        st.sidebar.success("Arquivos processados com sucesso!")
    except ValueError as e:
        st.sidebar.error(str(e))

# Criar armazenamento vetorial
vector_store = None
if documents:
    vector_store = create_vector_store(documents, api_key)
    st.success("Armazenamento vetorial criado com sucesso!")

# Inicializar o chatbot com o armazenamento vetorial
try:
    chatbot = ChatOpenAI(api_key=api_key, model="gpt-3.5-turbo", request_timeout=30)
except Exception as e:
    st.error(f"Erro ao conectar com o modelo OpenAI: {str(e)}")
    st.stop()

# Chatbot interativo
st.header("Chatbot Interativo")

# Inicializar o estado da sessão para rastrear envio de mensagens
if "chat_submitted" not in st.session_state:
    st.session_state.chat_submitted = False  # Indica se uma mensagem foi enviada
if "messages" not in st.session_state:
    st.session_state.messages = []

# Campo de entrada para o chat
user_input = st.text_input("Digite sua pergunta:")

# Botão explícito para enviar a mensagem
chat_submitted = st.button("Enviar Mensagem")

# Marcar o envio do chat apenas quando o botão for clicado
if chat_submitted and user_input:
    st.session_state.chat_submitted = True
    st.session_state.messages.append(HumanMessage(content=user_input))  # Adiciona a mensagem do usuário

# Processar a resposta apenas quando o botão for clicado
if st.session_state.chat_submitted:
    if vector_store:
        from langchain.chains import RetrievalQA
        retriever = vector_store.as_retriever()
        qa_chain = RetrievalQA.from_chain_type(
            llm=chatbot.model,
            retriever=retriever
        )
        # Obter resposta com base no vetor
        response = qa_chain.run(user_input)
    else:
        # Resposta padrão do chatbot
        response = chatbot.respond(st.session_state.messages)

    # Adicionar a resposta ao histórico
    st.session_state.messages.append(AIMessage(content=response))
    st.session_state.chat_submitted = False  # Resetar o estado após o processamento

# Exibir histórico de mensagens
for msg in st.session_state.messages:
    if isinstance(msg, SystemMessage):
        st.write(f"**System:** {msg.content}")
    elif isinstance(msg, HumanMessage):
        st.write(f"**User:** {msg.content}")
    elif isinstance(msg, AIMessage):
        st.write(f"**Assistant:** {msg.content}")

# Geração de Planos de Aula Personalizados
st.header("Criação de Planos de Aula Personalizados")

# Inicializa o estado da sessão para rastrear o envio
if "lesson_plan" not in st.session_state:
    st.session_state.lesson_plan = None
if "form_submitted" not in st.session_state:
    st.session_state.form_submitted = False

# Botão para resetar a memória e criar nova aula
if st.button("Gerar Nova Aula"):
    chatbot.reset_lesson_memory()
    st.session_state.lesson_plan = None
    st.session_state.form_submitted = False  # Reseta o estado do formulário
    st.success("Memória de aula resetada. Pronto para uma nova aula!")

# Formulário para entrada de dados
with st.form(key="lesson_plan_form"):
    topic = st.text_input("Tema da aula", placeholder="Exemplo: Direito Imobiliário")
    level = st.selectbox("Nível dos alunos", ["Fundamental", "Médio", "Superior"])
    duration = st.text_input("Duração da aula", placeholder="Exemplo: 1 hora")
    submitted = st.form_submit_button("Gerar Plano de Aula")  # Botão para envio do formulário

    # Marca o formulário como enviado apenas quando o botão for clicado
    if submitted:
        st.session_state.form_submitted = True

# Lógica de geração de plano de aula, acionada somente após envio
if st.session_state.form_submitted:
    if topic and level and duration:
        lesson_plan = chatbot.generate_lesson_plan(topic, level, duration)
        st.session_state.lesson_plan = lesson_plan
        st.subheader("Plano de Aula Gerado:")
        st.write(lesson_plan)
        st.session_state.form_submitted = False  # Evita re-execução ao alterar os campos
    else:
        st.warning("Por favor, preencha todos os campos antes de gerar o plano de aula.")
elif st.session_state.lesson_plan:
    # Exibe o plano de aula existente
    st.subheader("Plano de Aula Atual:")
    st.write(st.session_state.lesson_plan)

# Verificar se um plano de aula foi gerado
if st.session_state.lesson_plan:
    st.subheader("Plano de Aula Atual:")
    st.write(st.session_state.lesson_plan)
    
    # Botão para download do plano de aula
    docx_file = export_to_docx(st.session_state.lesson_plan)
    st.download_button(
        label="Baixar Plano de Aula em DOCX",
        data=docx_file,
        file_name="Plano_de_Aula.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


