import os
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.schema import Document
from docx import Document as DocxDocument
from tempfile import NamedTemporaryFile
import logging

logger = logging.getLogger(__name__)

def save_to_temp_file(uploaded_file):
    """Salva o arquivo enviado pelo Streamlit em um arquivo temporário."""
    temp_file = NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1])
    temp_file.write(uploaded_file.read())
    temp_file.close()
    logger.info(f"Arquivo salvo temporariamente: {temp_file.name}")
    return temp_file.name

def load_pdf(file_path):
    """Carrega e processa arquivos PDF."""
    logger.info(f"Processando arquivo PDF: {file_path}")
    loader = PyPDFLoader(file_path)
    return loader.load()

def load_docx(file_path):
    """Carrega e processa arquivos DOCX."""
    logger.info(f"Processando arquivo DOCX: {file_path}")
    doc = DocxDocument(file_path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text)]

def load_txt(file_path):
    """Carrega e processa arquivos TXT."""
    logger.info(f"Processando arquivo TXT: {file_path}")
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    return [Document(page_content=content)]

def process_uploaded_files(uploaded_files):
    """Processa os arquivos enviados pelo usuário e retorna uma lista de documentos."""
    documents = []
    for uploaded_file in uploaded_files:
        try:
            # Salva o arquivo temporariamente
            temp_file_path = save_to_temp_file(uploaded_file)

            # Identifica o tipo de arquivo e processa
            if uploaded_file.name.endswith(".pdf"):
                documents.extend(load_pdf(temp_file_path))
            elif uploaded_file.name.endswith(".docx"):
                documents.extend(load_docx(temp_file_path))
            elif uploaded_file.name.endswith(".txt"):
                documents.extend(load_txt(temp_file_path))
            else:
                raise ValueError(f"Formato de arquivo não suportado: {uploaded_file.name}")

        except Exception as e:
            logger.error(f"Erro ao processar o arquivo {uploaded_file.name}: {str(e)}", exc_info=True)
            raise RuntimeError(f"Erro ao processar o arquivo {uploaded_file.name}: {str(e)}") from e

        finally:
            # Remove o arquivo temporário, se existir
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                logger.info(f"Arquivo temporário removido: {temp_file_path}")

    return documents
